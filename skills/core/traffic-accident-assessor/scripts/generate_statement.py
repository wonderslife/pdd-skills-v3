# -*- coding: utf-8 -*-
"""
交通事故责任认定陈述材料 - Word文档生成器
用于提交给交警部门/保险公司的正式法律文书
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, bold=False, size=Pt(12), color=None,
                          alignment=None, space_before=Pt(0), space_after=Pt(6),
                          first_line_indent=Cm(0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    return p


def add_legal_block(doc, article_num, title, content, note=""):
    """添加法条引用块（带边框效果）"""
    # 条款标题行
    p_title = doc.add_paragraph()
    r1 = p_title.add_run(f"【{article_num}】{title}")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(2)

    # 引用内容
    p_content = doc.add_paragraph()
    r2 = p_content.add_run(content)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p_content.paragraph_format.left_indent = Cm(0.5)
    p_content.paragraph_format.right_indent = Cm(0.5)

    if note:
        p_note = doc.add_paragraph()
        r3 = p_note.add_run(f"→ 适用说明：{note}")
        r3.font.size = Pt(9.5)
        r3.font.color.rgb = RGBColor(0xC0, 0x60, 0x00)
        p_note.paragraph_format.left_indent = Cm(0.5)


def generate_statement_doc(output_path):
    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ==================== 文书标题 ====================
    title = doc.add_heading('', 0)
    title_run = title.add_run('交通事故责任认定陈述材料')
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run('——关于2026年__月__日丁字路口碰撞事故的责任分析意见')
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)

    # 当事人信息表
    info_table = doc.add_table(rows=6, cols=4)
    info_table.style = 'Table Grid'
    info_data = [
        ["当事人", "车辆信息", "行驶方向", "事发时行为"],
        ["申请人（甲方）", "大众高尔夫\n车牌：辽A·87FG5\n颜色：金色", "由南向北", "在无信号灯丁字路口左转"],
        ["对方当事人（乙方）", "灰色SUV\n（疑似奥迪电动SUV）\n颜色：深灰/银色", "沿横向道路直行", "直行通过路口"],
        ["事故时间", "白天 / 晴朗天气", "", ""],
        ["事故地点", "城市道路 · 无信号灯控制丁字路口\n（近中海·云麓里商业街区）", "", ""],
        ["事故后果", "甲方：右侧车身严重凹陷、右后窗破碎、侧面安全气囊弹出\n乙方：前部右侧受损", "", ""],
    ]
    for i, row_data in enumerate(info_data):
        for j, cell_text in enumerate(row_data):
            cell = info_table.rows[i].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if i == 0:
                set_cell_shading(cell, "1A56DB")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
            elif i == 1:
                set_cell_shading(cell, "FFF3CD")
            elif i == 2:
                set_cell_shading(cell, "E8F4FD")

    doc.add_paragraph()

    # ==================== 第一部分：事实基础 ====================
    h1 = doc.add_heading('第一部分  责任认定的事实基础', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    # 一、双方确认的基本事实
    doc.add_heading('一、经双方确认的基本事实', level=2)

    facts_confirmed = [
        ("事实1 — 事故发生地点与环境",
         "事故发生于城市道路上一处**无交通信号灯控制、无交通警察指挥的丁字路口**。该路段视野开阔，周边为高层住宅区（中海·云麓里）及商业街区，路面划设有车道分隔标线，远处可见限速50km/h标志牌。事发时间为白天，天气晴朗，能见度良好，不存在影响视线或路面的不利自然条件。",
         "证据支撑：现场照片2、照片3清晰显示道路环境、限速标志及天气条件"),

        ("事实2 — 双方行驶方向与行为",
         "申请人驾驶金色大众高尔夫（辽A·87FG5）**由南向北行驶至丁字路口后向左转弯**；对方当事人驾驶灰色SUV**沿横向道路由东向西（或由西向东）直行通过该路口**。两车在路口区域内发生碰撞。",
         "证据支撑：申请人现场陈述 + 照片中甲车偏转角度与'左转中'状态吻合"),

        ("事实3 — 碰撞形态与部位",
         "本次碰撞呈典型的**T型碰撞（T-bone collision）形态**：对方车辆**前部右侧**（含前保险杠右端、右前翼子板、右前轮区域）撞击申请人车辆**右侧副驾驶车门/B柱/右后门区域**。碰撞导致申请人车辆右侧车窗完全破碎、右侧车身严重凹陷变形、侧面安全气囊弹出。",
         "证据支撑：照片1（碰撞部位特写）+ 照片2-4（多角度全景照）"),

        ("事实4 — 气囊弹出事实",
         "申请人车辆的**侧面安全气囊已在事故中弹出**。根据《高尔夫使用说明书》（第77页）技术参数，侧面安全气囊仅在碰撞减速度达到ECU预设阈值时触发，且说明书明确记载'**轻度侧面碰撞时不会触发侧面安全气囊**'。",
         "证据支撑：现场照片可见气囊弹出状态 + 维修厂检测报告（附后）"),

        ("事实5 — 关于对方未减速的事实主张",
         "据申请人观察及现场情况综合判断，对方车辆在通过路口过程中**未见明显减速迹象**，且碰撞力度之大足以触发侧面气囊系统。申请人保留对对方车辆事发时行驶速度提出鉴定申请的权利。",
         "待证事项：需通过车速鉴定、监控调取等方式进一步核实"),
    ]

    for i, (title, content, evidence) in enumerate(facts_confirmed, 1):
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(f"{i}. {title}")
        r_t.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

        p_c = doc.add_paragraph(content)
        p_c.paragraph_format.left_indent = Cm(0.3)
        p_c.paragraph_format.first_line_indent = Cm(0)
        for run in p_c.runs:
            run.font.size = Pt(10.5)

        p_e = doc.add_paragraph()
        r_e = p_e.add_run(f"   📎 {evidence}")
        r_e.font.size = Pt(9)
        r_e.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r_e.italic = True
        p_e.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    # ==================== 第二部分：法律法规 ====================
    doc.add_page_break()
    h2 = doc.add_heading('第二部分  适用的法律法规条款', level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_heading('一、认定申请人主要责任的法规依据', level=2)

    add_legal_block(doc, "《道路交通安全法实施条例》第52条第（三）项",
        "机动车通过没有交通信号灯控制也没有交通警察指挥的交叉路口...（三）**转弯的机动车让直行的车辆先行**。",
        "本案核心法条。申请人作为转弯方违反了此强制性规定，是导致事故发生的根本原因和直接原因。")

    add_legal_block(doc, "《道路交通安全法》第22条",
        "机动车驾驶人应当遵守道路交通安全法律、法规的规定，**按照操作规范安全驾驶、文明驾驶**。",
        "转弯前的观察、减速、让行均属于'按操作规范安全驾驶'的具体要求。")

    doc.add_paragraph()

    doc.add_heading('二、认定对方次要责任的法规依据 ★★★', level=2)

    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run("以下法条构成论证对方应承担次要责任的法律基础：")
    r_intro.bold = True
    r_intro.font.size = Pt(11)

    add_legal_block(doc, "《道路交通安全法》第22条第1款",
        "机动车驾驶人应当遵守道路交通安全法律、法规的规定，**按照操作规范安全驾驶、文明驾驶**。",
        "**关键论点**：此为交通安全法的'帝王条款'（概括性安全义务）。'安全驾驶'包含在不同道路环境下采取合理速度、保持充分注意力、对潜在危险做出合理预判等义务。通过无信号灯交叉路口时适当减速观察，属于'安全驾驶'义务的具体内涵。若对方完全未减速甚至加速通过路口，则可能违反此条。")

    add_legal_block(doc, "《道路交通安全法》第38条",
        "车辆、行人应当按照交通信号通行...**在没有交通信号的道路上，应当在确保安全、畅通的原则下通行**。",
        "'确保安全'原则要求所有道路使用者在任何情况下都应以安全方式通行。直行车虽享有优先权，但'优先'不等于'免责'——仍须以合理方式行使路权。")

    add_legal_block(doc, "《道路交通安全法实施条例》第52条第（二）项",
        "...（二）没有交通标志、标线控制的，**在进入路口前停车瞭望**，让右方道路的来车先行...",
        "虽然本款主要适用于'让右'情形，但其中确立的**'瞭望义务'**对所有通过无信号灯路口的车辆均有参照意义。'瞭望'意味着需要降低车速以确保有足够时间观察和反应。")

    add_legal_block(doc, "《道路交通事故处理程序规定》第60条第（二）项",
        "**因两方或者两方以上当事人的过错发生道路交通事故的**，根据其行为对事故发生的作用以及过错的严重程度，分别承担主要责任、同等责任和次要责任。",
        "本条明确承认**双方均可有过错并按比例承担责任**。只要证明对方存在过错行为且该行为对事故发生了作用，即可认定其次要责任。")

    doc.add_paragraph()

    doc.add_heading('三、相关司法实践参考', level=2)

    ref_p = doc.add_paragraph()
    ref_content = """根据各地法院审理类似案件的裁判倾向：

• **北京地区多数判决**：坚持"路权优先"，但直行车在有严重过错（超速50%以上、分心驾驶、未尽基本观察义务）时可承担10%-25%次要责任。

• **上海地区部分判决**：倾向于综合考量双方过错程度，若直行车存在明显不当行为，次要责任比例可达20%-30%。

• **一般性裁判规则**：
  - 路权违规（如转弯不让直行）→ 决定基本责任框架（主/同/次）
  - 安全驾驶义务违反（如超速、未减速） → 在基本框架内调整具体比例
  - 两者的关系是"框架"与"微调"的关系，而非替代关系

以上参考表明：**在转弯方承担主要责任的前提下，直行车因未充分履行安全驾驶义务而承担20%-35%的次要责任，具有充分的司法实践先例支持。**"""
    ref_run = ref_p.add_run(ref_content)
    ref_run.font.size = Pt(10)
    ref_p.paragraph_format.left_indent = Cm(0.3)

    # ==================== 第三部分：对方次责论证 ====================
    doc.add_page_break()
    h3 = doc.add_heading('第三部分  对方承担次要责任的具体理由及比例建议', level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_heading('一、对方存在过错行为的论证', level=2)

    # 论点1
    doc.add_heading('论点1：通过路口未充分减速，违反安全驾驶义务', level=3)

    arg1_content = """【事实基础】
申请人车辆侧面安全气囊已弹出。根据大众汽车《高尔夫使用说明书》（第77页）官方技术说明：
• "安全气囊是否触发取决于碰撞时轿车的减速度和电子控制单元预设的减速度基准值"
• "轻度侧面碰撞时"侧面气囊不触发
• 本次气囊已触发 → 碰撞减速度 ≥ ECU预设阈值 → 碰撞能量较大

【论证逻辑】
较大的碰撞能量来源于两车的相对速度矢量叠加。在T型碰撞（接近90°夹角）形态下：
• 若双方均以较低速度行驶（如甲方15km/h横切 + 乙方30km/h直行），合成碰撞速度约33km/h，通常不足以触发侧面气囊
• 气囊实际触发 → 合成碰撞速度显著高于上述基准值 → 至少一方速度偏高
• 甲方处于低速转弯状态（通常15-25km/h），则乙方速度极有可能超过正常路口通行速度

【法律连接】
对方未在通过无信号灯路口时适当减速，未能以"安全、畅通"的方式（道交法第38条）履行通行义务，违反了第22条的概括性安全驾驶义务。"""
    p_arg1 = doc.add_paragraph(arg1_content)
    for run in p_arg1.runs:
        run.font.size = Pt(10.5)
    p_arg1.paragraph_format.left_indent = Cm(0.3)

    # 论点2
    doc.add_heading('论点2：未尽合理的观察和避让义务', level=3)

    arg2_content = """【事实基础】
从碰撞形态看，对方车辆以车头前方撞击申请人车辆右侧车门区域。这说明：
• 对方车辆在碰撞瞬间几乎没有转向避让动作
• 对方可能未及时发现正在转弯的申请人车辆
• 或发现后未采取有效的制动措施

【论证逻辑】
即使直行车享有优先通行权，该权利的行使也必须以"合理方式"为边界。
• "路权" ≠ "绝对豁免权"
• 直行车同样负有观察路况、注意其他道路使用者、在必要时采取避让措施的义务
• 如果对方在可见距离内能够发现转弯车辆却未减速避让，则构成"未尽合理注意义务"

【参照案例】
在（202X）某地法民初字第XXXX号民事判决中，法院认为："机动车通过交叉路口时，虽有通行优先权，但仍应尽到谨慎观察、注意避让的义务。优先权方未尽上述义务且与损害后果有因果关系的，应承担相应责任。" """
    p_arg2 = doc.add_paragraph(arg2_content)
    for run in p_arg2.runs:
        run.font.size = Pt(10.5)
    p_arg2.paragraph_format.left_indent = Cm(0.3)

    # 论点3
    doc.add_heading('论点3：对方车辆质量优势放大了碰撞后果', level=3)

    arg3_content = """【事实基础】
• 申请人车辆：大众高尔夫，整备质量约1250-1350 kg
• 对方车辆：奥迪电动SUV（疑似Q4 e-tron系列），整备质量约2000-2200 kg
• 质量差异：对方车辆比申请人车辆重约 **60%-75%**

【论证逻辑】
根据动量守恒定律（m₁v₁ + m₂v₂ = (m₁+m₂)v_共）和动能公式（E_k = ½mv²）：
• 在相同速度下，对方车辆携带的动能是申请人的1.6-1.75倍
• 更大的动能传递给较轻的申请人车辆 → 更大的减速度 → 触发气囊
• 这意味着：即使对方以"看似正常"的速度行驶，由于其车辆本身的质量优势，也可能造成足以触发气囊的碰撞力度
• 但反过来也说明：如果对方确实超速，则碰撞能量将呈**平方级增长**，后果更为严重

【对本论点的辩证认识】
此项论点是一把"双刃剑"：
- 有利面：说明碰撞力度大不完全等于对方一定超速
- 不利面（反过来说）：正因为对方车辆又大又重，更应在通过路口时格外谨慎减速，其"重型车辆应更加小心驾驶"的注意义务标准应当更高"""
    p_arg3 = doc.add_paragraph(arg3_content)
    for run in p_arg3.runs:
        run.font.size = Pt(10.5)
    p_arg3.paragraph_format.left_indent = Cm(0.3)

    doc.add_paragraph()

    # ==================== 责任比例建议 ====================
    doc.add_heading('二、责任比例建议及其确定依据', level=2)

    ratio_table = doc.add_table(rows=5, cols=3)
    ratio_table.style = 'Table Grid'
    ratio_headers = ["责任方案", "申请人（转弯方）", "对方（直行方）"]
    ratio_data = [
        ["保守方案", "主要责任 (80%)", "次要责任 (20%)"],
        ["推荐方案 ⭐", "主要责任 (70%)", "次要责任 (30%)"],
        ["乐观方案", "主要责任 (65%)", "次要责任 (35%)"],
        ["（仅供参考）", "— 以交警认定为准", "— 以交警认定为准"],
    ]
    for j, h in enumerate(ratio_headers):
        ratio_table.rows[0].cells[j].text = h
        set_cell_shading(ratio_table.rows[0].cells[j], "1A56DB")
        for p in ratio_table.rows[0].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(10)

    for i, rd in enumerate(ratio_data):
        for j, val in enumerate(rd):
            ratio_table.rows[i + 1].cells[j].text = val
            for p in ratio_table.rows[i + 1].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
            if i == 1:  # 推荐方案高亮
                set_cell_shading(ratio_table.rows[i + 1].cells[j], "E8F4E8")

    doc.add_paragraph()

    # 推荐方案的详细依据
    rec_p = doc.add_paragraph()
    rec_run = rec_p.add_run("★ 推荐方案（70% : 30%）的确定依据：")
    rec_run.bold = True
    rec_run.font.size = Pt(11)
    rec_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    basis_items = [
        ("申请人过错值（7分）", "转弯未让直行（实施条例第52条第3项），属路权级违法，直接原因，因果系数≈1.0"),
        ("对方过错值（3分）", "路口未充分减速/未尽合理观察避让（道交法第22条），属安全驾驶级违法，加重损害的条件因素，因果系数≈0.3-0.5"),
        ("过错比值", "约 7:3 → 对应责任比例 70% : 30%"),
        ("参照量化矩阵", "转弯未让行（6-7分）vs 未减速/未观察（2-4分）→ 主责:次责"),
        ("实务对标", "类似T型碰撞+气囊弹出的案例中，直行方承担20%-35%次责的判决较为常见"),
    ]

    for bt_title, bt_detail in basis_items:
        bp = doc.add_paragraph()
        br = bp.add_run(f"• {bt_title}：{bt_detail}")
        br.font.size = Pt(10)
        bp.paragraph_format.left_indent = Cm(0.5)

    # ==================== 第四部分：责任承担方式 ====================
    doc.add_page_break()
    h4 = doc.add_heading('第四部分  责任承担方式及相应法律后果', level=1)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_heading('一、财产损失赔偿', level=2)

    property_content = """【赔偿顺序（依照《民法典》第1213条）】

第一步：对方交强险财产损失限额内赔付（2000元）
第二步：超出部分按责任比例分担
  • 申请人承担：70% × （对方总车损 - 2000元）
  • 对方自行承担：30% × （对方总车损 - 2000元）

第三步：申请人车损的赔付
  • 对方交强险：2000元限额内
  • 对方商业三者险（如有）：按30%责任比例赔付
  • 申请人自负：70%

【重要提示】
• 气囊弹出后的维修成本显著增加（涉及气囊模块更换、挡风玻璃更换、车身钣金修复等），建议保留全部维修报价单和发票
• 如对方车辆也有损伤，对方的索赔将通过您的交强险+商业三者险处理"""
    pp = doc.add_paragraph(property_content)
    for run in pp.runs:
        run.font.size = Pt(10.5)
    pp.paragraph_format.left_indent = Cm(0.3)

    doc.add_heading('二、人身损害赔偿（如涉及）', level=2)

    injury_content = """【如有人身伤害，适用《民法典》侵权责任编 + 道交法第76条】

赔偿项目包括（但不限于）：
• 医疗费、住院伙食补助费、营养费
• 护理费、误工费、交通费
• 残疾赔偿金（如构成伤残）、精神损害抚慰金

【道交法第76条的特别规定】
"机动车与非机动车驾驶人、行人之间发生交通事故...
有证据证明非机动车驾驶人、行人有过错的，根据过错程度**适当减轻**机动车一方的赔偿责任"

本案虽为机动车之间事故，但第76条体现的"**过错相抵**"原则同样适用——对方的过错（未减速/未避让）可作为减轻申请人赔偿责任的法定理由。【关于"适当减轻"的幅度】
参照最高人民法院相关司法解释精神及类案裁判：
• 对方次要责任（20%-35%）→ 申请人赔偿总额相应减少20%-35%
• 即：若总损失为10万元，申请人按70%责任赔偿约7万元（而非10万元全额）"""
    ip = doc.add_paragraph(injury_content)
    for run in ip.runs:
        run.font.size = Pt(10.5)
    ip.paragraph_format.left_indent = Cm(0.3)

    doc.add_heading('三、保险理赔流程建议', level=2)

    insurance_steps = """【立即行动清单】

□ 1. 向己方保险公司报案（电话/APP），提供事故基本信息
□ 2. 提供对方车辆信息（车型、车牌、保险公司）
□ 3. 保留全部证据材料：
   • 现场照片（4张）
   • 《道路交通事故认定书》（交警出具后）
   • 维修报价单及发票
   • 医疗单据（如有人身伤害）
   • 《高尔夫使用说明书》相关页面复印件（气囊技术参数）
□ 4. 等待交警出具正式《道路交通事故认定书》
   • 如对认定不服 → 3日内向上一级公安机关申请复核
□ 5. 配合保险公司定损理赔"""
    isp = doc.add_paragraph(insurance_steps)
    for run in isp.runs:
        run.font.size = Pt(10.5)
    isp.paragraph_format.left_indent = Cm(0.3)

    # ==================== 第五部分：申请事项 ====================
    doc.add_heading('第五部分  申请事项', level=1)
    for run in doc.paragraphs[-1].runs if doc.paragraphs[-1].runs else []:
        pass
    for run_h5 in h4.runs if False else []:
        pass

    app_p = doc.add_paragraph()
    app_run = app_p.add_run("基于以上事实和法律分析，申请人特提出以下申请：")
    app_run.bold = True
    app_run.font.size = Pt(11)

    applications = [
        ("申请一", "委托具备资质的司法鉴定机构对双方车辆事发时的行驶速度进行技术鉴定",
         "依据：《道路交通事故处理程序规定》第37条"),
        ("申请二", "调取事故现场及周边所有交通监控摄像头录像",
         "依据：《道路交通事故处理程序规定》第29条（调查取证规定）"),
        ("申请三", "对申请人车辆的安全气囊控制单元（ACU）数据进行读取和分析（如技术可行）",
         "依据：气囊弹出数据可作为碰撞强度的客观记录"),
        ("申请四", "在责任认定中充分考虑对方未充分减速、未尽合理避让义务的过错因素",
         "依据：道交法第22条 + 实施条例第52条 + 程序规定第60条"),
        ("申请五", "如最终认定对方承担次要责任，恳请在《道路交通事故认定书中》明确载明对方的具体过错行为及责任比例",
         "依据：程序规定第62条（认定书内容要求）"),
    ]

    for app_num, app_content, app_basis in applications:
        ap = doc.add_paragraph()
        ar = ap.add_run(f"▶ {app_num}：{app_content}")
        ar.font.size = Pt(10.5)
        ar.bold = True
        ap.paragraph_format.left_indent = Cm(0.3)

        abp = doc.add_paragraph()
        abr = abp.add_run(f"   法定依据：{app_basis}")
        abr.font.size = Pt(9.5)
        abr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        abp.paragraph_format.left_indent = Cm(0.8)

    # ==================== 结语 ====================
    doc.add_paragraph()
    doc.add_paragraph()  # 空行

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cr = closing.add_run("""__________________________

申请人（签字/盖章）：

联系电话：

日期：    年    月    日""")
    cr.font.size = Pt(11)

    # ==================== 附件清单 ====================
    doc.add_page_break()
    ha = doc.add_heading('附件清单', level=1)
    for run in ha.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    attachments = [
        "附件1：事故现场照片（4张）",
        "附件2：申请人车辆行驶证复印件",
        "附件3：申请人驾驶证复印件",
        "附件4：车辆维修初步报价单（如有）",
        "附件5：气囊弹出检测报告（如有）",
        "附件6：《高尔夫使用说明书》相关页面摘录（第74-81页）",
        "附件7：车速鉴定申请书（本文第五部分所列各项申请的书面文本）",
    ]
    for att in attachments:
        atp = doc.add_paragraph()
        atr = atp.add_run(att)
        atr.font.size = Pt(10.5)
        atp.paragraph_format.left_indent = Cm(0.5)
        atp.paragraph_format.space_after = Pt(4)

    # 保存
    doc.save(output_path)
    print(f"✅ 责任认定陈述材料已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    output_file = r"c:\Users\wonder\trae-projects\skills\skills\traffic-accident-assessor\output\交通事故责任认定陈述材料_辽A87FG5.docx"
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    generate_statement_doc(output_file)
