# -*- coding: utf-8 -*-
"""
交通事故责任评估报告 - Word文档生成器
生成包含照片、专业排版、法律引用的完整评估报告
"""

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_with_style(doc, text, level=1, color=None):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading


def create_info_table(doc, data):
    """创建信息表格"""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        # 设置表头样式（第一行加粗背景）
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "E8F4FD")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
        else:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
    return table


def add_legal_quote(doc, article_num, title, content, note=""):
    """添加法律条文引用块"""
    # 条款标题
    p = doc.add_paragraph()
    run_title = p.add_run(f"◆ 第{article_num}条 【{title}】")
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run_title.font.size = Pt(11)

    # 引用内容（缩进+灰色背景效果通过边框模拟）
    p2 = doc.add_paragraph()
    run_content = p2.add_run(f'  "{content}"')
    run_content.italic = True
    run_content.font.size = Pt(10)
    run_content.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 适用说明
    if note:
        p3 = doc.add_paragraph()
        run_note = p3.add_run(f"  ▸ 适用说明：{note}")
        run_note.font.size = Pt(9.5)
        run_note.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_warning_box(doc, text):
    """添加警告/免责框"""
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {text}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xD4, 0x73, 0x13)


def generate_report(output_path, image_paths=None):
    """生成完整评估报告"""

    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ==================== 封面标题 ====================
    title = doc.add_heading('交通事故责任评估报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run('Traffic Accident Liability Assessment Report')
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run_sub.italic = True

    doc.add_paragraph()

    # 报告元信息表格
    meta_data = [
        ("评估编号", "TAA-20260404-001"),
        ("评估日期", datetime.now().strftime("%Y年%m月%d日")),
        ("事故类型", "机动车之间 · 丁字路口碰撞（转弯 vs 直行）"),
        ("事故地点", "城市道路 · 无信号灯丁字路口（近中海·云麓里）"),
        ("天气条件", "晴朗 · 能见度良好 · 白天"),
    ]
    create_info_table(doc, meta_data)
    doc.add_paragraph()

    # ==================== 一、事故经过概述 ====================
    doc.add_heading('一、事故经过概述', level=1)

    overview_text = """当事人甲（金色大众高尔夫，车牌号 辽A·87FG5）驾驶车辆由南向北行驶至一处无红绿灯控制的丁字路口时向左转弯。此时，当事人乙（灰色SUV，疑似奥迪电动SUV）沿横向道路直行通过该路口。两车在路口区域内发生碰撞，乙车前部撞击甲车右侧副驾驶位车门区域，造成甲车右后门车窗破碎、右侧车身严重凹陷，气囊弹出。"""
    p = doc.add_paragraph(overview_text)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.5

    # ==================== 二、现场照片分析 ====================
    doc.add_heading('二、现场照片分析', level=1)

    photo_descriptions = [
        {
            "title": "照片1 — 碰撞部位特写",
            "desc": """关键发现：
• 甲车损伤位置：右侧车身（副驾驶门/B柱/右后门区域），右后侧窗完全破碎
• 乙车损伤位置：前部右侧（前保险杠右端/右前翼子板/右前轮区域）
• 碰撞形态：T型碰撞（T-bone），乙车前右 → 甲车右侧
• 判定意义：确认甲车处于横穿状态（侧面受撞），符合"转弯车切入直行车路径"的典型碰撞模式""",
        },
        {
            "title": "照片2 — 全景照（角度一）",
            "desc": """关键发现：
• 两车最终位置：甲车呈明显偏转角度（约30°-45°），与"左转中"状态高度吻合
• 限速标志：远处可见限速50km/h标志牌
• 道路环境：多车道城市道路，高层住宅区背景（中海·云麓里）
• 天气条件：晴天，阳光强烈，视线条件极佳，排除天气因素影响""",
        },
        {
            "title": "照片3 — 全景照（角度二）",
            "desc": """关键发现：
• 甲车车牌清晰可辨：辽A·87FG5
• 路面标线：可见白色车道分隔线，属划设标线的正规道路
• 周边动态：有骑行电动车人员、远处其他车辆，该路口有一定交通流量
• 路口形态：开阔T型/十字路口区域，视距充足""",
        },
        {
            "title": "照片4 — 后方视角",
            "desc": """关键发现：
• 甲车尾部完整可见，车牌辽A·87FG5二次确认
• 乙车位于甲车左后方，两车呈夹角停靠（与碰撞后自然滑移规律一致）
• 周边环境：商业街区（绿源电动车店、蓝色经典等商铺）
• 目击人员：现场有多名围观群众，可能存在目击者""",
        },
    ]

    for i, info in enumerate(photo_descriptions):
        # 子标题
        h = doc.add_heading(info["title"], level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

        # 插入图片（如果有）
        if image_paths and i < len(image_paths) and os.path.exists(image_paths[i]):
            img_path = image_paths[i]
            try:
                doc.add_picture(img_path, width=Inches(5.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = doc.add_paragraph(f"图{i+1}: {info['title']}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    run.italic = True
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {e}]")
        else:
            # 占位提示
            p_placeholder = doc.add_paragraph()
            p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ph = p_placeholder.add_run(f"[{info['title']} — 请手动插入对应照片]")
            run_ph.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            run_ph.italic = True

        # 分析描述
        desc_p = doc.add_paragraph(info["desc"])
        desc_p.paragraph_format.left_indent = Cm(0.5)

        doc.add_paragraph()

    # 综合结论框
    doc.add_heading('📌 照片证据链综合结论', level=2)
    conclusion_p = doc.add_paragraph()
    conclusion_p.paragraph_format.left_indent = Cm(0.5)
    conclusion_p.paragraph_format.right_indent = Cm(0.5)
    run_c = conclusion_p.add_run(
        '四张照片共同构成的证据链一致指向同一事实 —— '
        '您的车辆在左转过程中被直行车辆从侧面撞击。'
        '碰撞形态（T型）、车辆最终位置（偏转角度）、受损部位（右侧车门区域）'
        '三者相互印证，完全符合"转弯车未让直行车"的事故类型特征。'
    )
    run_c.font.size = Pt(10.5)
    run_c.bold = True

    # ASCII态势图
    doc.add_paragraph()
    diagram = doc.add_paragraph()
    diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_diag = diagram.add_run("""
    ┌─────────────────────────┐
    │         【北】           │
    │          ↑              │
    │          │ 甲车(金)      │
    │          │  ↖ 左转中     │
    │          │   ╲ ← 碰撞点  │
    └─────┬────┼──────┼────────┘
          │    │      │
     ←────┼── 乙车(灰) ──→
     (直行)   (直行)
        [西]    [东]

    ★ T型碰撞：转弯车切入直行车路径 ★
    """)
    run_diag.font.name = "Consolas"
    run_diag.font.size = Pt(9)

    # ==================== 三、当事人行为分析 ====================
    doc.add_heading('三、当事人行为分析', level=1)

    # 当事人A表格
    doc.add_heading('▶ 当事人A（您方 — 金色大众高尔夫 · 左转）', level=2)
    table_a = doc.add_table(rows=6, cols=2)
    table_a.style = 'Table Grid'
    a_data = [
        ("行为描述", "由南向北行驶至无信号灯丁字路口时执行左转操作，横穿横向道路"),
        ("是否存在违法", "✅ 是 — 违反转弯车辆让行直行车辆的法定义务"),
        ("违反法条", "《道路交通安全法实施条例》第52条第（三）项"),
        ("过错程度", "高 — 转弯让直行是基本路权规则，未履行即构成对直行方路权的侵犯"),
        ("因果测试", "若甲方不左转或待直行车辆通过后再转，事故不会发生 → 直接原因"),
        ("过错分值参考", "6-7分（未按规定让行，参照过错量化矩阵）"),
    ]
    for i, (k, v) in enumerate(a_data):
        table_a.rows[i].cells[0].text = k
        table_a.rows[i].cells[1].text = v
        if i == 0:
            for cell in table_a.rows[i].cells:
                set_cell_shading(cell, "FFF3CD")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)

    doc.add_paragraph()

    # 当事人B表格
    doc.add_heading('▶ 当事人B（对方 — 灰色SUV · 直行）', level=2)
    table_b = doc.add_table(rows=6, cols=2)
    table_b.style = 'Table Grid'
    b_data = [
        ("行为描述", "沿横向道路直行通过丁字路口，与正在左转的甲车发生碰撞"),
        ("是否存在违法", "⚠️ 初步判断无明显违法 — 直行车辆依法享有优先通行权"),
        ("潜在审查事项", "①是否超过50km/h限速；②是否采取合理避让措施；③是否存在分心驾驶"),
        ("过错程度", "低或无 — 直行方在无信号灯路口处于法律保护的优先地位"),
        ("因果测试", "即使换为其他直行车辆，若甲方仍违规转弯，事故仍大概率发生 → 非必要原因"),
        ("过错分值参考", "0-2分（如无超速等则为0分）"),
    ]
    for i, (k, v) in enumerate(b_data):
        table_b.rows[i].cells[0].text = k
        table_b.rows[i].cells[1].text = v
        if i == 0:
            for cell in table_b.rows[i].cells:
                set_cell_shading(cell, "D4EDDA")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)

    # ==================== 四、责任判定结论 ====================
    doc.add_page_break()
    doc.add_heading('四、责任判定结论', level=1)

    # 判定结果总表
    result_table = doc.add_table(rows=3, cols=4)
    result_table.style = 'Table Grid'
    headers = ["当事人", "车辆信息", "责任等级", "责任比例"]
    for j, h in enumerate(headers):
        result_table.rows[0].cells[j].text = h
        set_cell_shading(result_table.rows[0].cells[j], "1A56DB")
        for p in result_table.rows[0].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)

    # A方数据
    result_table.rows[1].cells[0].text = "当事人A（您方）"
    result_table.rows[1].cells[1].text = "金色大众高尔夫\n辽A·87FG5"
    result_table.rows[1].cells[2].text = "主要责任"
    result_table.rows[1].cells[3].text = "70% ~ 90%"
    for cell in result_table.rows[1].cells:
        set_cell_shading(cell, "FFE5E5")

    # B方数据
    result_table.rows[2].cells[0].text = "当事人B（对方）"
    result_table.rows[2].cells[1].text = "灰色SUV"
    result_table.rows[2].cells[2].text = "无责 ~ 次要"
    result_table.rows[2].cells[3].text = "0% ~ 30%"
    for cell in result_table.rows[2].cells:
        set_cell_shading(cell, "E5F5E5")

    doc.add_paragraph()

    # 判定推理过程
    doc.add_heading('判定推理过程', level=2)

    steps = [
        ("第一步 — 违法行为认定", "A方违反《实施条例》第52条第3项（转弯让直行）；B方初步审查无违法。"),
        ("第二步 — 因果关系分析（But-for测试）", "若A方不左转→事故不发生（直接原因，系数≈1.0）；换任何直行车辆→事故仍发生（B方非原因）。"),
        ("第三步 — 过错程度比较", "A方路权弱势、有明确让行义务、行为危险性高（6-7分）；B方路权强势、一般安全义务即可（0-2分）。过错比 A:B ≈ 7:1 ~ 9:1。"),
        ("第四步 — 综合评定", "第52条第3项确立刚性规则 → 转弯方负有法定让行义务 → 未履行 → 进入对方合法路径 → 无法避免碰撞 → 承担主要以上责任。"),
    ]

    for step_title, step_content in steps:
        sp = doc.add_paragraph()
        run_st = sp.add_run(f"● {step_title}\n  {step_content}")
        run_st.font.size = Pt(10.5)
        sp.paragraph_format.left_indent = Cm(0.3)
        sp.paragraph_format.space_after = Pt(8)

    # ==================== 五、法律依据详列 ====================
    doc.add_page_break()
    doc.add_heading('五、法律依据详列', level=1)

    # 道交法条文
    doc.add_heading('《中华人民共和国道路交通安全法》', level=2)

    add_legal_quote(doc, "38", "信号灯及通行原则",
        "车辆、行人应当按照交通信号通行；遇有交通警察现场指挥时，应当按照交通警察的指挥通行；在没有交通信号的道路上，应当在确保安全、畅通的原则下通行。",
        "无信号灯路口的基本通行原则——转弯车辆在未确保安全的情况下转弯即违反此条。")

    doc.add_paragraph()

    add_legal_quote(doc, "22", "安全驾驶义务",
        "机动车驾驶人应当遵守道路交通安全法律、法规的规定，按照操作规范安全驾驶、文明驾驶。",
        "转弯前的观察、减速、让行均属于'按操作规范安全驾驶'的具体要求。")

    doc.add_paragraph()

    # 实施条例（核心！）
    doc.add_heading('《中华人民共和国道路交通安全法实施条例》★★★ 核心法条 ★★★', level=2)

    add_legal_quote(doc, "52", "无信号灯路口通行规则",
        "机动车通过没有交通信号灯控制也没有交通警察指挥的交叉路口，除应当遵守第五十一条第（二）项、第（三）项的规定外，还应当遵守下列规定：\n（一）有交通标志、标线控制的，让优先通行的一方先行；\n（二）没有交通标志、标线控制的，在进入路口前停车瞭望，让右方道路的来车先行；\n（三）转弯的机动车让直行的车辆先行；\n（四）相对方向行驶的右转弯的机动车让左转弯的车辆先行。",
        "⭐ 第（三）项是本案的核心法律依据！'转弯的机动车让直行的车辆先行'是强制性规定，不是建议。违反此规定即构成违法，且通常与事故的发生具有直接因果关系。")

    doc.add_paragraph()

    add_legal_quote(doc, "51条第(七)项", "转弯让直行（有信号灯路口参照）",
        "...（七）在没有方向指示信号灯的交叉路口，转弯的机动车让直行的车辆、行人通行。相对方向行驶的右转弯的机动车让左转弯的车辆先行。",
        "体现立法者对'转弯让直行'规则的一贯态度，即使在有信号灯路口同样适用。")

    doc.add_paragraph()

    # 程序规定
    doc.add_heading('《道路交通事故处理程序规定》（2020年修正）', level=2)

    add_legal_quote(doc, "60", "责任认定标准",
        "公安机关交通管理部门经过调查后，应当根据当事人的行为对发生道路交通事故所起的作用以及过错的严重程度，确定当事人的责任：（一）因一方当事人的过错导致道路交通事故的，承担全部责任；（二）因两方或者两方以上当事人的过错发生道路交通事故的，根据其行为对事故发生的作用以及过错的严重程度，分别承担主要责任、同等责任和次要责任...",
        "本案适用第（二）项：A方过错起主要作用且程度较高 → 对应'主要责任'档次。")

    # ==================== 六、法律建议 ====================
    doc.add_page_break()
    doc.add_heading('六、法律建议', level=1)

    # 您方的建议
    doc.add_heading('👤 对您方（转弯方 / 主要责任方）的建议', level=2)
    your_advice = [
        "配合交警调查，如实陈述事实 —— 不要隐瞒或歪曲事故经过，照片证据已清楚反映碰撞形态",
        "主动争取减轻责任的因素 —— 如已提前开转向灯（≥3秒）、已部分进入路口后才被撞、对方明显超速等",
        "立即联系保险公司报案 —— 车损由对方交强险先赔2000元，超出部分按责任比例分担；对方的损失由您的交强险+三者险赔付",
        "保留所有单据 —— 气囊弹出意味着维修成本增加，如有人身伤害务必保留全部医疗单据",
        "如对认定书不服 —— 可在收到认定书之日起3日内申请复核",
    ]
    for i, advice in enumerate(your_advice, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {advice}")
        run.font.size = Pt(10.5)
        p.paragraph_format.left_indent = Cm(0.3)

    doc.add_paragraph()

    # 对方的建议
    doc.add_heading('👤 对对方（直行方）的建议', level=2)
    their_advice = [
        "主张无责论据 —— 强调正常直行中被转弯车辆切入路径，提供照片证明T型碰撞形态",
        "如有行车记录仪视频 —— 这是最有力的证据",
        "防范对方抗辩 —— 对方可能主张'已打转向灯'或'已转了一半'，但这些不能免除让行义务",
    ]
    for i, advice in enumerate(their_advice, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {advice}")
        run.font.size = Pt(10.5)
        p.paragraph_format.left_indent = Cm(0.3)

    # 后续流程
    doc.add_heading('📋 后续处理流程', level=2)
    flow = doc.add_paragraph()
    flow_text = """
    ① 交警现场勘查 → 等待《道路交通事故认定书》
       ↓
    ② 收到认定书 → 如不服可在3日内申请复核
       ↓
    ③ 责任确定后 → 双方保险公司介入理赔
       ↓
    ④ 协商赔偿 → 协不成则诉讼解决
    """
    run_flow = flow.add_run(flow_text)
    run_flow.font.size = Pt(10.5)
    flow.paragraph_format.left_indent = Cm(1)

    # ==================== 七、免责声明 ====================
    doc.add_page_break()
    doc.add_heading('七、免责声明', level=1)

    disclaimer_items = [
        "本评估报告基于用户提供的文字描述和4张现场照片生成，仅供参考之用。",
        "本报告不能替代公安机关交通管理部门出具的正式《道路交通事故认定书》，最终责任认定以交警部门的官方调查结论为准。",
        "本分析未考虑以下可能影响判定的因素：是否有交通监控录像、是否有目击者证言、双方驾驶员是否存在酒驾/毒驾/无证驾驶、路面是否有特殊标志标线、具体车速鉴定结果等。",
        "如对方车辆存在超速（超过50km/h限速20%以上）、分心驾驶或未尽合理观察义务等情况，责任比例可能向转弯方有利方向调整。",
        "如涉及人身伤害或较大财产损失，建议及时委托专业律师处理后续理赔事宜。",
        "本报告中引用的法律条文均为公开有效的法律法规内容。",
    ]

    for item in disclaimer_items:
        dp = doc.add_paragraph()
        run_dp = dp.add_run(f"⚠ {item}")
        run_dp.font.size = Pt(10)
        run_dp.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # 页脚
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_footer = footer_p.add_run(f"— — —\n生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}\nTraffic Accident Assessor Skill v1.0")
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run_footer.italic = True

    # 保存
    doc.save(output_path)
    print(f"✅ 报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    output_file = r"c:\Users\wonder\trae-projects\skills\skills\traffic-accident-assessor\output\交通事故责任评估报告_辽A87FG5.docx"

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    # 图片路径（如有实际图片则传入）
    image_paths = [
        r"c:\Users\wonder\trae-projects\skills\skills\traffic-accident-assessor\output\photo1.jpg",
        r"c:\Users\wonder\trae-projects\skills\skills\traffic-accident-assessor\output\photo2.jpg",
        r"c:\Users\wonder\trae-projects\skills\skills\traffic-accident-assessor\output\photo3.jpg",
        r"c:\Users\wonder\trae-projects\skills\skills\traffic-accident-assessor\output\photo4.jpg",
    ]

    generate_report(output_file, image_paths)
