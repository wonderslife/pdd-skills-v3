# -*- coding: utf-8 -*-
"""
交通事故责任认定陈述材料 - GB/T 9704-2012公文格式版生成器
基于 official-doc-writer 技能的党政机关公文格式标准改造
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml


# ============================================================
# GB/T 9704-2012 公文格式常量定义
# ============================================================

# 页面参数
PAGE_WIDTH = Cm(21)       # A4宽度
PAGE_HEIGHT = Cm(29.7)    # A4高度
MARGIN_TOP = Cm(3.7)      # 天头 37mm ± 1mm
MARGIN_BOTTOM = Cm(3.5)   # 地脚
MARGIN_LEFT = Cm(2.8)     # 订口 28mm ± 1mm
MARGIN_RIGHT = Cm(2.6)    # 翻口

# 字号对照表（磅值）
FONT_SIZE_ER_HAO = Pt(22)      # 二号 — 标题
FONT_SIZE_SAN_HAO = Pt(16)     # 三号 — 正文/一级标题
FONT_SIZE_SI_HAO = Pt(14)      # 四号 — 抄送等
FONT_SIZE_SI_BAN = Pt(12)      # 四号半 — 页码

# 字体名称（GB/T 9704-2012 规定字体）
FONT_TITLE = "方正小标宋简体"     # 标题：小标宋体
FONT_BODY = "仿宋"               # 正文：仿宋_GB2312
FONT_HEADING1 = "黑体"           # 一级标题：黑体
FONT_HEADING2 = "楷体"           # 二级标题：楷体_GB2312
FONT_HEADING3 = "仿宋"           # 三级标题：仿宋
FONT_COPY = "仿宋"                # 抄送/印发：仿宋
FONT_PAGE_NUM = "宋体"            # 页码：宋体

# 行距参数
LINE_SPACING_PT = 28.0            # 固定值28磅 ≈ 单倍行距（每页22行）
CHARS_PER_LINE = 28              # 每行28字


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name, font_size, bold=False, color=None, italic=False):
    """统一设置run的字体属性"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_paragraph_format(para, line_spacing=LINE_SPACING_PT,
                         first_line_indent=None, space_before=Pt(0),
                         space_after=Pt(0), alignment=None,
                         left_indent=None, right_indent=None):
    """设置段落格式（符合公文规范）"""
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    else:
        # 正文默认首行缩进2字符（约2cm）
        para.paragraph_format.first_line_indent = Cm(0.85)
    if left_indent is not None:
        para.paragraph_format.left_indent = left_indent
    if right_indent is not None:
        para.paragraph_format.right_indent = right_indent
    para.paragraph_format.space_before = space_before
    para.paragraph_format.space_after = space_after
    if alignment:
        para.alignment = alignment


def add_doc_title(doc, text):
    """添加公文式标题（二号小标宋，居中，红色分隔线下空二行）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(20), space_after=Pt(16))
    run = p.add_run(text)
    set_run_font(run, FONT_TITLE, FONT_SIZE_ER_HAO, bold=True,
                 color=RGBColor(0xFF, 0x00, 0x00))  # 红色标题
    return p


def add_subtitle(doc, text):
    """添加副标题"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(0), space_after=Pt(6))
    run = p.add_run(text)
    set_run_font(run, FONT_BODY, Pt(14), color=RGBColor(0x66, 0x66, 0x66))  # 宋体副标题
    return p


def add_heading_level1(doc, text):
    # 一级标题：黑体三号，编号"一、"
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0),
                         space_before=Pt(12), space_after=Pt(4))
    run = p.add_run(text)
    set_run_font(run, FONT_HEADING1, FONT_SIZE_SAN_HAO, bold=True)
    return p


def add_heading_level2(doc, text):
    # 二级标题：楷体三号，编号"（一）"
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0),
                         space_before=Pt(8), space_after=Pt(3))
    run = p.add_run(text)
    set_run_font(run, FONT_HEADING2, FONT_SIZE_SAN_HAO)
    return p


def add_heading_level3(doc, text):
    # 三级标题：仿宋三号加粗，编号"1."
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0),
                         space_before=Pt(6), space_after=Pt(2))
    run = p.add_run(text)
    set_run_font(run, FONT_HEADING3, FONT_SIZE_SAN_HAO, bold=True)
    return p


def add_body_text(doc, text, indent=True, bold_parts=None):
    """正文段落：仿宋三号"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0.85) if indent else Cm(0),
                         space_before=Pt(0), space_after=Pt(3))

    if bold_parts:
        # 支持部分文字加粗
        remaining = text
        for bp in bold_parts:
            if bp in remaining:
                parts = remaining.split(bp, 1)
                if parts[0]:
                    r1 = p.add_run(parts[0])
                    set_run_font(r1, FONT_BODY, FONT_SIZE_SAN_HAO)
                r_bold = p.add_run(bp)
                set_run_font(r_bold, FONT_BODY, FONT_SIZE_SAN_HAO, bold=True)
                remaining = parts[1] if len(parts) > 1 else ""
        if remaining:
            r_last = p.add_run(remaining)
            set_run_font(r_last, FONT_BODY, FONT_SIZE_SAN_HAO)
    else:
        run = p.add_run(text)
        set_run_font(run, FONT_BODY, FONT_SIZE_SAN_HAO)
    return p


def add_legal_quote_block(doc, article_num, title, content, note=""):
    """法条引用块（带缩进和特殊格式）"""
    # 条款标题
    p_title = doc.add_paragraph()
    set_paragraph_format(p_title, first_line_indent=Cm(0.5),
                         space_before=Pt(8), space_after=Pt(2))
    r1 = p_title.add_run(f"【{article_num}】{title}")
    set_run_font(r1, FONT_HEADING1, Pt(13), bold=True,
                 color=RGBColor(0x1A, 0x56, 0xDB))

    # 引用内容（缩进+斜体）
    p_content = doc.add_paragraph()
    set_paragraph_format(p_content, first_line_indent=Cm(0.85),
                         left_indent=Cm(0.5), right_indent=Cm(0.5),
                         space_before=Pt(0), space_after=Pt(2))
    r2 = p_content.add_run(content)
    set_run_font(r2, FONT_BODY, Pt(13), italic=True)

    # 适用说明
    if note:
        p_note = doc.add_paragraph()
        set_paragraph_format(p_note, first_line_indent=Cm(0.5),
                             left_indent=Cm(0.5), space_before=Pt(0), space_after=Pt(6))
        r3 = p_note.add_run(f"→ 适用说明：{note}")
        set_run_font(r3, FONT_BODY, Pt(11.5), color=RGBColor(0xC0, 0x60, 0x00))


def add_info_table(doc, data, header_color="1A56DB"):
    """创建信息表格（公文风格）"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, FONT_BODY, Pt(10.5))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                set_cell_shading(cell, header_color)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
    return table


def add_separator_line(doc, color="FF0000", width_pt=0.75):
    """添加红色分隔线"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0),
                         space_before=Pt(4), space_after=Pt(4))
    # 底部边框模拟分隔线
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(width_pt * 8)))  # 转换为八分之一磅
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_official_statement(output_path):
    """生成 GB/T 9704-2012 格式的责任认定陈述材料"""

    doc = Document()

    # ========== 页面设置（GB/T 9704-2012）==========
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_SAN_HAO
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # ==================== 版头区域 ====================

    # 发文机关标志（红色）
    header_p = doc.add_paragraph()
    set_paragraph_format(header_p, first_line_indent=Cm(0),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(10), space_after=Pt(4))
    hr = header_p.add_run("交通事故当事人陈述材料")
    set_run_font(hr, FONT_TITLE, Pt(18), color=RGBColor(0xCC, 0x00, 0x00))

    # 红色分隔线
    add_separator_line(doc, "CC0000", 1.0)

    # ==================== 标题 ====================
    add_doc_title(doc, "关于2026年__月__日丁字路口碰撞事故")
    add_doc_title(doc, "的责任认定陈述意见")

    # 副标题
    add_subtitle(doc, "——申请人：辽A·87FG5 驾驶人")

    # 空行
    doc.add_paragraph()

    # ==================== 主送机关 ====================
    recip_p = doc.add_paragraph()
    set_paragraph_format(recip_p, first_line_indent=Cm(0),
                         space_before=Pt(0), space_after=Pt(6))
    rr = recip_p.add_run("___________公安局交通警察支队：")
    set_run_font(rr, FONT_BODY, FONT_SIZE_SAN_HAO, bold=True)

    # ==================== 正文开头 ====================
    opening = (
        "本人系2026年__月__日在____路口发生的道路交通事故的当事人之一"
        "（驾驶辽A·87FG5号大众高尔夫轿车）。"
        "现就本次事故的责任认定问题，根据《中华人民共和国道路交通安全法》"
        "及相关法律法规之规定，提出如下陈述意见。"
    )
    add_body_text(doc, opening)

    # ==================== 第一部分 ====================
    add_heading_level1(doc, "一、事故基本事实")

    # 当事人信息表格
    info_data = [
        ["项目", "申请人（甲方）", "对方当事人（乙方）"],
        ["车辆信息", "大众高尔夫\n车牌：辽A·87FG5\n颜色：金色", "灰色SUV\n（疑似奥迪电动SUV）\n颜色：深灰"],
        ["行驶方向", "由南向北行驶后左转", "沿横向道路直行"],
        ["事发行为", "在无信号灯丁字路口左转", "直行通过路口"],
        ["车辆受损", "右侧车身凹陷、右后窗破碎\n侧面安全气囊弹出", "前部右侧受损"],
    ]
    add_info_table(doc, info_data)
    doc.add_paragraph()

    add_body_text(doc, "经现场确认，事故发生于城市道路上的一处无交通信号灯控制、")
    add_body_text(doc, "亦无交通警察指挥的丁字路口。事发时间为白天，天气晴朗，")
    add_body_text(doc, "能见度良好。该路段视野开阔，路面划设有车道分隔标线，")
    add_body_text(doc, "远处可见限速50km/h标志牌。两车在路口区域内发生碰撞，")
    add_body_text(doc, "呈典型T型碰撞形态（对方车前部撞击申请人车辆右侧车门区域）。")

    # 关键事实分项
    add_heading_level2(doc, "（一）关键事实确认")

    facts = [
        ("1.", "碰撞形态事实",
         "对方车辆前部右侧撞击申请人车辆右侧副驾驶位车门/B柱/右后门区域，"
         "造成右后侧窗完全破碎、车身严重凹陷、侧面安全气囊弹出。"
         "上述事实有现场照片4张为证。"),
        ("2.", "气囊弹出事实",
         "申请人车辆的侧面安全气囊已在事故中触发并弹出。"
         "根据大众汽车《高尔夫使用说明书》（第77页）技术规定："
         '"安全气囊是否触发取决于碰撞时轿车的减速度和电子控制单元预设的'
         '减速度基准值"，且明确记载"轻度侧面碰撞时不会触发侧面安全气囊"。'
         "本次气囊已触发的事实表明碰撞能量较大。"),
        ("3.", "关于对方车速的主张",
         "据申请人观察及综合判断，对方车辆在通过路口过程中未见明显减速迹象。"
         "鉴于碰撞力度足以触发精密的安全气囊系统，"
         "申请人保留对对方事发时行驶速度提出司法鉴定的权利。"),
    ]

    for num, title, content in facts:
        hp3 = doc.add_paragraph()
        set_paragraph_format(hp3, first_line_indent=Cm(0),
                             space_before=Pt(6), space_after=Pt(1))
        hr_num = hp3.add_run(f"{num}{title}")
        set_run_font(hr_num, FONT_BODY, FONT_SIZE_SAN_HAO, bold=True)

        cp = doc.add_paragraph()
        set_paragraph_format(cp, space_before=Pt(0), space_after=Pt(4))
        cr = cp.add_run(content)
        set_run_font(cr, FONT_BODY, FONT_SIZE_SAN_HAO)

    # ==================== 第二部分 ====================
    add_heading_level1(doc, "二、法律法规依据")

    add_heading_level2(doc, "（一）认定申请人主要责任的法规")

    add_legal_quote_block(doc,
        "《道路交通安全法实施条例》第52条第（三）项",
        "机动车通过没有交通信号灯控制也没有交通警察指挥的交叉路口... "
        "（三）转弯的机动车让直行的车辆先行。",
        "本案核心法条。申请人作为转弯方违反此强制性规定，是导致事故发生的根本原因。")

    add_legal_quote_block(doc,
        "《道路交通安全法》第22条第1款",
        "机动车驾驶人应当遵守道路交通安全法律、法规的规定，"
        "按照操作规范安全驾驶、文明驾驶。",
        "转弯前的观察、减速、让行均属于'按操作规范安全驾驶'的具体要求。")

    add_heading_level2(doc, "（二）认定对方次要责任的法规 ★★★")

    add_body_text(doc, "以下法条构成论证对方应承担次要责任的法律基础：")

    add_legal_quote_block(doc,
        "《道路交通安全法》第22条第1款",
        "机动车驾驶人应当遵守道路交通安全法律、法规的规定，"
        "按照操作规范安全驾驶、文明驾驶。",
        "此为交通安全法的'帝王条款'。'安全驾驶'包含在不同道路环境下采取合理速度、"
        "保持充分注意力、对潜在危险做出合理预判等义务。"
        "通过无信号灯交叉路口时适当减速观察，属于'安全驾驶'义务的具体内涵。")

    add_legal_quote_block(doc,
        "《道路交通安全法》第38条",
        "车辆、行人应当按照交通信号通行；遇有交通警察现场指挥时，应当按照"
        "交通警察的指挥通行；在没有交通信号的道路上，应当在确保安全、畅通的原则下通行。",
        "'确保安全'原则要求所有道路使用者以安全方式通行。"
        "直行车虽享有优先权，但'优先'不等于'免责'——仍须以合理方式行使路权。")

    add_legal_quote_block(doc,
        "《道路交通安全法实施条例》第52条第（二）项",
        "...（二）没有交通标志、标线控制的，在进入路口前停车瞭望，"
        "让右方道路的来车先行...",
        "其中确立的'瞭望义务'对所有通过无信号灯路口的车辆均有参照意义。"
        "'瞭望'意味着需要降低车速以确保有足够时间观察和反应。")

    add_legal_quote_block(doc,
        "《道路交通事故处理程序规定》第60条第（二）项",
        "因两方或者两方以上当事人的过错发生道路交通事故的，"
        "根据其行为对事故发生的作用以及过错的严重程度，"
        "分别承担主要责任、同等责任和次要责任。",
        "本条明确承认双方均可有过错并按比例承担责任。"
        "只要证明对方存在过错且该行为与事故有关联，即可认定其次要责任。")

    # ==================== 第三部分 ====================
    add_heading_level1(doc, "三、对方应承担次要责任的论证")

    add_heading_level2(doc, "（一）论点一：未充分减速，违反安全驾驶义务")

    arg1 = (
        "申请人车辆侧面安全气囊已弹出这一客观事实，是本论点的核心技术支撑。"
        "根据大众汽车官方说明书的技术参数：（1）气囊触发取决于碰撞减速度是否达到"
        "ECU预设阈值；（2）轻度侧面碰撞不触发侧面气囊；（3）无法确定单一触发车速，"
        "但触发的必要条件是足够的碰撞能量。T型碰撞中，合成碰撞速度等于双方速度矢量之和。"
        "申请人处于低速转弯状态（通常15-25km/h），若对方也以正常谨慎速度（如20-30km/h）通过，"
        "则合成碰撞速度通常不足以触发侧面气囊。气囊实际触发的事实强烈暗示——"
        "至少一方速度偏高，而对方作为直行车且驾驶质量更大的SUV，"
        "更有可能是速度偏高的一方。"
    )
    add_body_text(doc, arg1)

    add_heading_level2(doc, "（二）论点二：未尽合理观察和避让义务")

    arg2 = (
        "从碰撞形态分析，对方车辆以车头前方撞击申请人右侧车门区域，"
        "说明碰撞瞬间对方几乎没有转向避让或有效制动。即使直行车依法享有优先通行权，"
        "该权利的行使必须以'合理方式'为边界。'路权'不等于'绝对豁免权'。"
        "直行车同样负有观察路况、注意其他道路使用者、必要时采取避让措施的义务。"
        "如果对方在可见距离内能够发现正在转弯的申请人车辆却未减速避让，"
        "即构成'未尽合理注意义务'。参照多地法院类案裁判，"
        "优先权方未尽注意义务且与损害后果有因果关系的，应承担相应比例责任。"
    )
    add_body_text(doc, arg2)

    add_heading_level2(doc, "（三）论点三：重型车辆应负有更高注意义务")

    arg3 = (
        "经查，申请人驾驶的大众高尔夫整备质量约1250-1350kg，"
        "对方驾驶的电动SUV整备质量约2000-2200kg，质量差异达60%-75%。"
        "根据动量守恒定律，在相同速度下对方车辆携带的动能是申请人的1.6至1.75倍。"
        "更重的车辆在碰撞中将传递更大的能量给较轻的一方。"
        "因此，驾驶重型车辆通过潜在冲突区域时，理应以更高的谨慎标准履行注意义务。"
        "对方未能因其车辆的质量优势而额外谨慎，亦属过失范畴。"
    )
    add_body_text(doc, arg3)

    # ==================== 第四部分：责任比例建议 ====================
    add_heading_level1(doc, "四、责任比例建议")

    add_heading_level2(doc, "（一）推荐方案")

    ratio_data = [
        ["方案类型", "申请人（甲方）", "对方（乙方）", "适用条件"],
        ["保守方案", "主要责任（80%）", "次要责任（20%）", "对方仅轻微过失"],
        ["推荐方案 ★", "主要责任（70%）", "次要责任（30%）", "对方存在明显未减速/未避让情形"],
        ["乐观方案", "主要责任（65%）", "次要责任（35%）", "对方被鉴定为超速或有其他严重过失"],
    ]
    t_ratio = add_info_table(doc, ratio_data, "2E7D32")
    doc.add_paragraph()

    add_heading_level2(doc, "（二）推荐方案的确定依据")

    basis_items = [
        "申请人过错值约7分：违反《实施条例》第52条第3项（转弯让直行），"
        "属路权级违法，直接原因，因果系数≈1.0。",
        "对方过错值约3分：违反《道交法》第22条（安全驾驶概括义务），"
        "属安全驾驶级违法，加重损害的条件因素，因果系数≈0.3-0.5。",
        "过错比值约为7:3，对应责任比例70%:30%。",
        "参照各地法院类似T型碰撞案件的裁判倾向，"
        "直行车因未充分履行安全驾驶义务而承担20%-35%次责具有充分的先例支持。",
    ]
    for bi in basis_items:
        bp = doc.add_paragraph()
        set_paragraph_format(bp, space_before=Pt(0), space_after=Pt(2))
        br = bp.add_run(f"• {bi}")
        set_run_font(br, FONT_BODY, FONT_SIZE_SAN_HAO)

    # ==================== 第五部分：申请事项 ====================
    add_heading_level1(doc, "五、正式申请事项")

    applications = [
        ("（一）", "申请委托具备资质的司法鉴定机构对双方车辆事发时的行驶速度进行技术鉴定。",
         "依据：《道路交通事故处理程序规定》第37条"),
        ("（二）", "申请调取事故现场及周边所有交通监控摄像头录像。",
         "依据：《道路交通事故处理程序规定》第29条"),
        ("（三）", "申请读取申请人车辆安全气囊控制单元数据（如技术可行）。",
         "依据：气囊弹出数据可作为碰撞强度的客观记录"),
        ("（四）", "请求在责任认定中充分考虑对方未充分减速、未尽合理避让义务的过错因素。",
         "依据：《道交法》第22条 + 《实施条例》第52条 + 《程序规定》第60条"),
        ("（五）", "如最终认定对方承担次要责任，恳请在《道路交通事故认定书中》"
         "明确载明对方的具体过错行为及责任比例。",
         "依据：《程序规定》第62条"),
    ]

    for num, app_content, app_basis in applications:
        ap = doc.add_paragraph()
        set_paragraph_format(ap, first_line_indent=Cm(0),
                             space_before=Pt(6), space_after=Pt(1))
        ar = ap.add_run(f"{num}{app_content}")
        set_run_font(ar, FONT_BODY, FONT_SIZE_SAN_HAO, bold=True)

        abp = doc.add_paragraph()
        set_paragraph_format(abp, left_indent=Cm(0.8),
                             space_before=Pt(0), space_after=Pt(4))
        abr = abp.add_run(f"法定依据：{app_basis}")
        set_run_font(abr, FONT_BODY, Pt(11), color=RGBColor(0x66, 0x66, 0x66))

    # ==================== 结语 ====================
    doc.add_paragraph()  # 空行

    closing_text = (
        "以上陈述意见，恳请贵队在作出事故责任认定时予以充分考虑。"
        "申请人愿意配合贵队的进一步调查工作，并提供一切必要的协助。\n"
    )
    add_body_text(doc, closing_text)

    closing_end = doc.add_paragraph()
    set_paragraph_format(closing_end, first_line_indent=Cm(0),
                         space_before=Pt(16), space_after=Pt(0))
    cer = closing_end.add_run("特此陈述。")
    set_run_font(cer, FONT_BODY, FONT_SIZE_SAN_HAO, bold=True)

    # 署名区
    doc.add_paragraph()  # 空
    doc.add_paragraph()  # 空

    sign_p = doc.add_paragraph()
    set_paragraph_format(sign_p, first_line_indent=Cm(0),
                         alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         space_before=Pt(0), space_after=Pt(0))
    sr = sign_p.add_run("\n\n\n申请人（签字）：________________")
    set_run_font(sr, FONT_BODY, FONT_SIZE_SAN_HAO)

    contact_p = doc.add_paragraph()
    set_paragraph_format(contact_p, first_line_indent=Cm(0),
                         alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         space_before=Pt(8), space_after=Pt(0))
    cr = contact_p.add_run("联系电话：________________")
    set_run_font(cr, FONT_BODY, FONT_SIZE_SAN_HAO)

    date_p = doc.add_paragraph()
    set_paragraph_format(date_p, first_line_indent=Cm(0),
                         alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         space_before=Pt(8), space_after=Pt(0))
    dr = date_p.add_run("日期：    年    月    日")
    set_run_font(dr, FONT_BODY, FONT_SIZE_SAN_HAO)

    # ==================== 版记区域 ====================
    doc.add_page_break()

    # 版记分隔线（粗线）
    add_separator_line(doc, "000000", 0.75)

    # 附件清单
    att_h = doc.add_paragraph()
    set_paragraph_format(att_h, first_line_indent=Cm(0),
                         space_before=Pt(8), space_after=Pt(6))
    atr = att_h.add_run("附件")
    set_run_font(atr, FONT_HEADING1, FONT_SIZE_SAN_HAO, bold=True)

    attachments = [
        "1. 事故现场照片（4张）",
        "2. 申请人车辆行驶证复印件",
        "3. 申请人驾驶证复印件",
        "4. 《高尔夫使用说明书》相关页面摘录（第74-81页）",
        "5. 车辆维修初步报价单（如有）",
        "6. 气囊弹出检测报告（如有）",
    ]

    for att in attachments:
        atp = doc.add_paragraph()
        set_paragraph_format(atp, first_line_indent=Cm(0),
                             left_indent=Cm(0.5), space_before=Pt(0), space_after=Pt(2))
        atrun = atp.add_run(att)
        set_run_font(atrun, FONT_BODY, FONT_SIZE_SAN_HAO)

    # 版记下分隔线（细线）
    add_separator_line(doc, "000000", 0.5)

    # 印发机关和日期
    print_p = doc.add_paragraph()
    set_paragraph_format(print_p, first_line_indent=Cm(0),
                         space_before=Pt(6), space_after=Pt(0))
    pr = print_p.add_run("（共7页）")
    set_run_font(pr, FONT_BODY, FONT_SIZE_SI_HAO, color=RGBColor(0x88, 0x88, 0x88))

    # 保存
    doc.save(output_path)
    print(f"\n{'='*60}")
    print(f"✅ GB/T 9704-2012 格式化陈述材料已生成")
    print(f"📄 文件路径: {output_path}")
    print(f"{'='*60}")
    return output_path


if __name__ == "__main__":
    output_file = r"c:\Users\wonder\trae-projects\skills\skills\traffic-accident-assessor\output\交通事故责任认定陈述材料_辽A87FG5_公文格式版.docx"
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    generate_official_statement(output_file)
