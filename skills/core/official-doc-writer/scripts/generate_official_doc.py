# -*- coding: utf-8 -*-
"""
党政机关公文生成脚本
符合GB/T 9704-2012《党政机关公文格式》国家标准
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def add_horizontal_line(doc, color='FF0000', size_pt=2):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run('_' * 100)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def create_official_document(doc_type, content):
    """
    创建符合GB/T 9704-2012标准的党政机关公文Word文档
    
    参数:
        doc_type: 公文类型（通知/报告/请示/函/通报/纪要等）
        content: 公文内容字典，包含以下字段：
            - issuer: 发文机关
            - doc_number: 发文字号
            - title: 标题
            - recipient: 主送机关
            - body: 正文内容（列表）
            - signer: 发文机关署名
            - date: 成文日期
            - attachment: 附件说明（可选）
            - copy_to: 抄送机关（可选）
            - issuer_office: 印发机关（可选，默认为发文机关办公室）
            - issue_date: 印发日期（可选，默认为成文日期）
    
    返回:
        Document对象
    """
    doc = Document()
    
    # 设置页面格式（A4，页边距符合GB/T 9704-2012标准）
    section = doc.sections[0]
    section.page_width = Cm(21)      # A4宽度
    section.page_height = Cm(29.7)   # A4高度
    section.top_margin = Cm(3.7)     # 天头37mm
    section.bottom_margin = Cm(3.0)  # 下白边
    section.left_margin = Cm(2.8)    # 订口28mm
    section.right_margin = Cm(2.6)   # 切口
    
    # ========== 版头部分 ==========
    
    # 发文机关标志（红色小标宋体）
    if content.get('issuer'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(content['issuer'] + '文件')
        run.font.size = Pt(22)  # 小标宋体
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        run.font.name = '方正小标宋_GBK'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋_GBK')
    
    # 发文字号（仿宋体3号）
    if content.get('doc_number'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(content['doc_number'])
        run.font.size = Pt(16)  # 3号字
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 红色分隔线
    add_horizontal_line(doc, color='FF0000', size_pt=2)
    
    # ========== 主体部分 ==========
    
    # 标题（小标宋体2号）
    if content.get('title'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(content['title'])
        run.font.size = Pt(22)  # 2号字
        run.font.bold = True
        run.font.name = '方正小标宋_GBK'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋_GBK')
    
    # 主送机关
    if content.get('recipient'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        run = p.add_run(content['recipient'] + '：')
        run.font.size = Pt(16)  # 3号字
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 正文内容
    if content.get('body'):
        for para_text in content['body']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(32)  # 首行缩进2字符
            run = p.add_run(para_text)
            run.font.size = Pt(16)  # 3号字
            run.font.name = '仿宋_GB2312'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 结尾语（根据公文类型）
    closing_phrases = {
        '通知': '特此通知。',
        '报告': '特此报告。',
        '请示': '妥否，请批示。',
        '函': '请予研究函复。',
        '通报': '特此通报。'
    }
    
    if doc_type in closing_phrases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(32)
        run = p.add_run(closing_phrases[doc_type])
        run.font.size = Pt(16)
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 附件说明
    if content.get('attachment'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(32)
        run = p.add_run('附件：' + content['attachment'])
        run.font.size = Pt(16)
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 发文机关署名
    if content.get('signer'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(content['signer'])
        run.font.size = Pt(16)
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 成文日期
    if content.get('date'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(content['date'])
        run.font.size = Pt(16)
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # ========== 版记部分 ==========
    
    # 版记分隔线（粗线）
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('═' * 100)
    run.font.size = Pt(3)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = '宋体'
    
    # 抄送机关
    if content.get('copy_to'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        run = p.add_run('抄送：' + content['copy_to'] + '。')
        run.font.size = Pt(14)  # 4号字
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 版记分隔线（细线）
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('─' * 100)
    run.font.size = Pt(2)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = '宋体'
    
    # 印发机关和印发日期
    issuer_office = content.get('issuer_office', content.get('issuer', '') + '办公室')
    issue_date = content.get('issue_date', content.get('date', ''))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f'{issuer_office}                    {issue_date}印发')
    run.font.size = Pt(14)  # 4号字
    run.font.name = '仿宋_GB2312'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    
    # 版记分隔线（粗线）
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('═' * 100)
    run.font.size = Pt(3)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = '宋体'
    
    return doc


if __name__ == '__main__':
    # 示例：生成报告类公文
    content = {
        'issuer': 'XXX公司',
        'doc_number': 'XX〔2026〕1号',
        'title': '关于XXX公司情况的报告',
        'recipient': 'XX市相关主管部门、XX投资集团',
        'body': [
            'XXX公司成立于2023年3月，注册资本5亿元，是XX投资集团下属企业，专注于数据相关业务。',
            '作为区域数字经济发展的主力军，公司定位清晰、职能明确，致力于推动本地数字经济发展。'
        ],
        'signer': 'XXX公司',
        'date': '2026年3月13日',
        'attachment': 'XXX公司核心成果及重点项目清单',
        'copy_to': 'XX市数字经济发展局、XX区人民政府'
    }
    
    doc = create_official_document('报告', content)
    doc.save('output.docx')
    print("公文已成功生成：output.docx")
