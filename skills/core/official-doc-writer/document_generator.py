# -*- coding: utf-8 -*-
"""
党政机关公文生成器
符合 GB/T 9704-2012《党政机关公文格式》国家标准
"""

import sys
import io
from datetime import datetime
from typing import Dict, List, Optional

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


class OfficialDocumentGenerator:
    """党政机关公文生成器"""
    
    FONT_SONGTI = '宋体'
    FONT_FANGSONG = '仿宋_GB2312'
    FONT_HEITI = '黑体'
    FONT_KAITI = '楷体_GB2312'
    FONT_XIAOBIAOSONG = '方正小标宋_GBK'
    
    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()
    
    def _setup_page(self):
        """设置页面格式"""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    
    def _setup_styles(self):
        """设置文档样式"""
        style = self.doc.styles['Normal']
        style.font.name = self.FONT_FANGSONG
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_FANGSONG)
        style.font.size = Pt(16)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    def _set_run_font(self, run, font_name: str, font_size: int, bold: bool = False, color: str = '000000'):
        """设置运行字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color != '000000':
            run.font.color.rgb = RGBColor.from_string(color)
    
    def _add_paragraph_with_font(self, text: str, font_name: str, font_size: int, 
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT, bold: bool = False,
                                   first_line_indent: int = 0, color: str = '000000',
                                   space_before: int = 0, space_after: int = 0):
        """添加带格式的段落"""
        p = self.doc.add_paragraph()
        p.alignment = alignment
        
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Pt(first_line_indent)
        
        if space_before > 0:
            p.paragraph_format.space_before = Pt(space_before)
        
        if space_after > 0:
            p.paragraph_format.space_after = Pt(space_after)
        
        if text:
            run = p.add_run(text)
            self._set_run_font(run, font_name, font_size, bold, color)
        
        return p
    
    def add_doc_classification(self, classification: str, period: str = None):
        """添加密级和保密期限"""
        text = classification
        if period:
            text += '★' + period
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_HEITI, 16, False)
        
        return p
    
    def add_urgency(self, urgency: str):
        """添加紧急程度"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = p.add_run(urgency)
        self._set_run_font(run, self.FONT_HEITI, 16, False)
        
        return p
    
    def add_document_number(self, doc_number: str):
        """添加发文字号"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(16)
        
        run = p.add_run(doc_number)
        self._set_run_font(run, self.FONT_FANGSONG, 16, False)
        
        return p
    
    def add_issuer_mark(self, issuer: str, is_red: bool = True):
        """添加发文机关标志"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(16)
        
        if '文件' not in issuer:
            issuer += '文件'
        
        run = p.add_run(issuer)
        color = 'FF0000' if is_red else '000000'
        self._set_run_font(run, self.FONT_XIAOBIAOSONG, 26, True, color)
        
        return p
    
    def add_signer(self, signer_name: str):
        """添加签发人（上行文）"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run1 = p.add_run('签发人：')
        self._set_run_font(run1, self.FONT_FANGSONG, 16, False)
        
        run2 = p.add_run(signer_name)
        self._set_run_font(run2, self.FONT_KAITI, 16, False)
        
        return p
    
    def add_red_separator(self):
        """添加红色分隔线"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(16)
        
        run = p.add_run('━' * 40)
        self._set_run_font(run, self.FONT_SONGTI, 16, False, 'FF0000')
        
        return p
    
    def add_title(self, title: str):
        """添加公文标题"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(16)
        
        run = p.add_run(title)
        self._set_run_font(run, self.FONT_XIAOBIAOSONG, 22, True)
        
        return p
    
    def add_recipient(self, recipient: str):
        """添加主送机关"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        
        run = p.add_run(recipient + '：')
        self._set_run_font(run, self.FONT_FANGSONG, 16, False)
        
        return p
    
    def add_body_paragraph(self, text: str, first_line_indent: int = 32):
        """添加正文段落"""
        p = self.doc.add_paragraph()
        
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Pt(first_line_indent)
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_FANGSONG, 16, False)
        
        return p
    
    def add_heading_level1(self, text: str):
        """添加一级标题（一、）"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        p.paragraph_format.space_before = Pt(8)
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_HEITI, 16, False)
        
        return p
    
    def add_heading_level2(self, text: str):
        """添加二级标题（（一））"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_KAITI, 16, False)
        
        return p
    
    def add_heading_level3(self, text: str):
        """添加三级标题（1.）"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_FANGSONG, 16, False)
        
        return p
    
    def add_attachment_note(self, attachments: List[str]):
        """添加附件说明"""
        if not attachments:
            return
        
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.first_line_indent = Pt(32)
        
        run = p.add_run('附件：')
        self._set_run_font(run, self.FONT_FANGSONG, 16, False)
        
        if len(attachments) == 1:
            run2 = p.add_run(attachments[0])
            self._set_run_font(run2, self.FONT_FANGSONG, 16, False)
        else:
            for i, att in enumerate(attachments, 1):
                p2 = self.doc.add_paragraph()
                p2.paragraph_format.first_line_indent = Pt(32)
                run2 = p2.add_run(f'{i}. {att}')
                self._set_run_font(run2, self.FONT_FANGSONG, 16, False)
    
    def add_issuer_signature(self, issuer: str, date: str = None):
        """添加发文机关署名和成文日期"""
        self.doc.add_paragraph()
        
        p1 = self.doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.right_indent = Pt(64)
        run1 = p1.add_run(issuer)
        self._set_run_font(run1, self.FONT_FANGSONG, 16, False)
        
        if date:
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.paragraph_format.right_indent = Pt(64)
            run2 = p2.add_run(date)
            self._set_run_font(run2, self.FONT_FANGSONG, 16, False)
    
    def add_note(self, note: str):
        """添加附注"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        
        run = p.add_run(f'（{note}）')
        self._set_run_font(run, self.FONT_FANGSONG, 16, False)
        
        return p
    
    def add_copy_send(self, copy_to: str):
        """添加抄送机关"""
        self._add_separator_line()
        
        p = self.doc.add_paragraph()
        run1 = p.add_run('抄送：')
        self._set_run_font(run1, self.FONT_FANGSONG, 14, False)
        
        run2 = p.add_run(copy_to + '。')
        self._set_run_font(run2, self.FONT_FANGSONG, 14, False)
        
        return p
    
    def add_print_info(self, print_org: str, print_date: str):
        """添加印发机关和印发日期"""
        p = self.doc.add_paragraph()
        
        run1 = p.add_run(print_org)
        self._set_run_font(run1, self.FONT_FANGSONG, 14, False)
        
        run2 = p.add_run(' ' * 20)
        self._set_run_font(run2, self.FONT_FANGSONG, 14, False)
        
        run3 = p.add_run(f'{print_date}印发')
        self._set_run_font(run3, self.FONT_FANGSONG, 14, False)
        
        self._add_separator_line()
    
    def _add_separator_line(self):
        """添加分隔线"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        run = p.add_run('─' * 50)
        self._set_run_font(run, self.FONT_SONGTI, 8, False)
    
    def add_closing(self, closing: str = '特此通知。'):
        """添加结尾语"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        
        run = p.add_run(closing)
        self._set_run_font(run, self.FONT_FANGSONG, 16, False)
        
        return p
    
    def save(self, filepath: str):
        """保存文档"""
        self.doc.save(filepath)
        return filepath


def create_notice(content: Dict) -> OfficialDocumentGenerator:
    """创建通知类公文"""
    gen = OfficialDocumentGenerator()
    
    if content.get('classification'):
        gen.add_doc_classification(content['classification'], content.get('classification_period'))
    
    if content.get('urgency'):
        gen.add_urgency(content['urgency'])
    
    gen.add_issuer_mark(content.get('issuer', ''))
    gen.add_document_number(content.get('doc_number', ''))
    
    if content.get('signer'):
        gen.add_signer(content['signer'])
    
    gen.add_red_separator()
    
    gen.add_title(content.get('title', ''))
    gen.add_recipient(content.get('recipient', ''))
    
    for para in content.get('body', []):
        if para.startswith('一、') or para.startswith('二、') or para.startswith('三、'):
            gen.add_heading_level1(para)
        elif para.startswith('（一）') or para.startswith('（二）'):
            gen.add_heading_level2(para)
        elif para.startswith('1.') or para.startswith('2.') or para.startswith('3.'):
            gen.add_heading_level3(para)
        else:
            gen.add_body_paragraph(para)
    
    if content.get('closing'):
        gen.add_closing(content['closing'])
    
    if content.get('attachments'):
        gen.add_attachment_note(content['attachments'])
    
    gen.add_issuer_signature(
        content.get('issuer_signature', content.get('issuer', '')),
        content.get('date', '')
    )
    
    if content.get('note'):
        gen.add_note(content['note'])
    
    if content.get('copy_to'):
        gen.add_copy_send(content['copy_to'])
    
    if content.get('print_org'):
        gen.add_print_info(
            content.get('print_org', ''),
            content.get('print_date', '')
        )
    
    return gen


def create_report(content: Dict) -> OfficialDocumentGenerator:
    """创建报告类公文"""
    gen = OfficialDocumentGenerator()
    
    gen.add_issuer_mark(content.get('issuer', ''))
    gen.add_document_number(content.get('doc_number', ''))
    gen.add_red_separator()
    
    gen.add_title(content.get('title', ''))
    gen.add_recipient(content.get('recipient', ''))
    
    for para in content.get('body', []):
        if para.startswith('一、') or para.startswith('二、') or para.startswith('三、'):
            gen.add_heading_level1(para)
        elif para.startswith('（一）') or para.startswith('（二）'):
            gen.add_heading_level2(para)
        else:
            gen.add_body_paragraph(para)
    
    gen.add_closing('特此报告。')
    
    gen.add_issuer_signature(
        content.get('issuer_signature', content.get('issuer', '')),
        content.get('date', '')
    )
    
    if content.get('copy_to'):
        gen.add_copy_send(content['copy_to'])
    
    return gen


def create_request(content: Dict) -> OfficialDocumentGenerator:
    """创建请示类公文"""
    gen = OfficialDocumentGenerator()
    
    gen.add_issuer_mark(content.get('issuer', ''))
    gen.add_document_number(content.get('doc_number', ''))
    
    if content.get('signer'):
        gen.add_signer(content['signer'])
    
    gen.add_red_separator()
    
    gen.add_title(content.get('title', ''))
    gen.add_recipient(content.get('recipient', ''))
    
    for para in content.get('body', []):
        if para.startswith('一、') or para.startswith('二、') or para.startswith('三、'):
            gen.add_heading_level1(para)
        elif para.startswith('（一）') or para.startswith('（二）'):
            gen.add_heading_level2(para)
        else:
            gen.add_body_paragraph(para)
    
    gen.add_closing('妥否，请批示。')
    
    if content.get('attachments'):
        gen.add_attachment_note(content['attachments'])
    
    gen.add_issuer_signature(
        content.get('issuer_signature', content.get('issuer', '')),
        content.get('date', '')
    )
    
    if content.get('note'):
        gen.add_note(content['note'])
    
    return gen


def create_letter(content: Dict) -> OfficialDocumentGenerator:
    """创建函类公文"""
    gen = OfficialDocumentGenerator()
    
    gen.add_issuer_mark(content.get('issuer', ''))
    gen.add_document_number(content.get('doc_number', ''))
    gen.add_red_separator()
    
    gen.add_title(content.get('title', ''))
    gen.add_recipient(content.get('recipient', ''))
    
    for para in content.get('body', []):
        gen.add_body_paragraph(para)
    
    gen.add_closing('请予研究函复。')
    
    gen.add_issuer_signature(
        content.get('issuer_signature', content.get('issuer', '')),
        content.get('date', '')
    )
    
    if content.get('note'):
        gen.add_note(content['note'])
    
    return gen


def create_minutes(content: Dict) -> OfficialDocumentGenerator:
    """创建纪要类公文"""
    gen = OfficialDocumentGenerator()
    
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content.get('meeting_name', '') + '纪要')
    gen._set_run_font(run, gen.FONT_XIAOBIAOSONG, 26, True, 'FF0000')
    
    gen.add_body_paragraph(content.get('time_location', ''))
    gen.add_body_paragraph(content.get('overview', ''))
    
    if content.get('attendees'):
        p = gen.doc.add_paragraph()
        run1 = p.add_run('出席：')
        gen._set_run_font(run1, gen.FONT_HEITI, 16, False)
        run2 = p.add_run(content['attendees'])
        gen._set_run_font(run2, gen.FONT_FANGSONG, 16, False)
    
    if content.get('absent'):
        p = gen.doc.add_paragraph()
        run1 = p.add_run('请假：')
        gen._set_run_font(run1, gen.FONT_HEITI, 16, False)
        run2 = p.add_run(content['absent'])
        gen._set_run_font(run2, gen.FONT_FANGSONG, 16, False)
    
    for para in content.get('body', []):
        if para.startswith('一、') or para.startswith('二、'):
            gen.add_heading_level1(para)
        else:
            gen.add_body_paragraph(para)
    
    return gen


DOCUMENT_TYPES = {
    '通知': create_notice,
    '报告': create_report,
    '请示': create_request,
    '函': create_letter,
    '纪要': create_minutes,
}


def generate_document(doc_type: str, content: Dict, output_path: str) -> str:
    """
    生成党政机关公文
    
    参数:
        doc_type: 公文类型（通知/报告/请示/函/纪要）
        content: 公文内容字典
        output_path: 输出文件路径
    
    返回:
        生成的文件路径
    """
    if doc_type not in DOCUMENT_TYPES:
        raise ValueError(f"不支持的公文类型: {doc_type}。支持的类型: {list(DOCUMENT_TYPES.keys())}")
    
    generator = DOCUMENT_TYPES[doc_type](content)
    generator.save(output_path)
    
    return output_path


if __name__ == '__main__':
    example_content = {
        'issuer': 'XXX公司',
        'doc_number': 'XX〔2024〕1号',
        'title': '关于开展互联网服务统一管理工作的通知',
        'recipient': '各部门、各子公司',
        'body': [
            '为进一步规范集团互联网服务管理，加强网络安全防护，提升信息化管理水平，根据国家网络安全相关法律法规及集团信息化建设总体规划要求，集团决定对各部门、各子公司在互联网上提供的服务实施统一管理。现将有关事项通知如下：',
            '一、工作目标',
            '通过对集团各部门、各子公司互联网服务资源的全面梳理与统一管理，建立健全互联网服务管理体系，消除安全隐患，保障信息系统安全稳定运行，提升集团整体信息化管理效能。',
            '二、主要内容',
            '请各部门、各子公司配合提供以下互联网服务相关信息：',
            '（一）微信公众号',
            '包括公众号名称、账号主体、运营负责人及联系方式。',
            '（二）网站信息',
            '包括网站名称、域名、IP地址、备案号、服务器位置。',
            '三、工作要求',
            '各单位要高度重视，认真组织，确保信息收集的全面性、准确性和及时性。',
        ],
        'closing': '特此通知。',
        'attachments': ['互联网服务信息登记表'],
        'issuer_signature': 'XXX公司',
        'date': '2024年1月15日',
        'note': '联系人：张三，联系电话：024-12345678',
        'copy_to': '集团各部门',
        'print_org': '集团办公室',
        'print_date': '2024年1月16日'
    }
    
    output = generate_document('通知', example_content, 'example_notice.docx')
    print(f"公文已生成: {output}")
